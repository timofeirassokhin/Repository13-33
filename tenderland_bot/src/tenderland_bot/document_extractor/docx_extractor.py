"""DOCX → text + tables через python-docx.

NB: это отличается от ``analyzer.extractor.extract_from_docx`` (тот возвращает
ExtractedSpec'и с КТРУ-логикой). Здесь — generic raw extraction для
**любого** DOCX (брошюры, наши шаблоны ТЗ, любые документы).

Для tender-specific extraction → analyzer.extractor.extract_from_docx.
"""
from __future__ import annotations

import logging
from pathlib import Path

from .core import ExtractedTable, ExtractionOutput

log = logging.getLogger(__name__)


def extract_from_docx_file(path: Path) -> ExtractionOutput:
    """Generic DOCX text+tables extraction."""
    out = ExtractionOutput(source_file=path, file_type="docx", extractor_name="python-docx")

    try:
        from docx import Document
    except ImportError:
        out.error = "python-docx not installed"
        return out

    try:
        doc = Document(str(path))
    except Exception as exc:
        out.error = f"failed to open DOCX: {exc}"
        return out

    # Paragraphs
    out.paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    out.text = "\n\n".join(out.paragraphs)

    # Tables
    for tbl in doc.tables:
        rows: list[list[str]] = []
        for row in tbl.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            out.tables.append(ExtractedTable(rows=rows))

    # Metadata
    try:
        cp = doc.core_properties
        out.metadata.update({
            "author": cp.author or "",
            "title": cp.title or "",
            "subject": cp.subject or "",
            "created": str(cp.created) if cp.created else "",
            "modified": str(cp.modified) if cp.modified else "",
        })
    except Exception:
        pass

    return out
