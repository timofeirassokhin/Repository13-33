"""Извлечение характеристик товара из ТЗ.

Стратегии (по приоритету для одного DOCX):

1. **КТРУ-таблица ФЗ-44** — типовая для государственных закупок.
   Колонки: «Наименование характеристики», «Значение характеристики»,
   «Единица измерения характеристики», «Тип характеристики», и т.п.
   Идентифицируется по headers первых 3 строк.

2. **Простая key-value таблица** — для коммерческих площадок.
   2 колонки: характеристика → значение. Иногда 3 (+ единица).

3. **Параграфный fallback** — `Характеристика: значение` / `Характеристика — значение`
   в обычных параграфах. Применяется когда таблиц нет или они мусорные.

Результат — `list[ExtractedSpec]`. Каждый spec нормализован минимально:
`name`, `value_raw` (текст из ячейки), `unit`, `type` (количеств./качеств./Bool),
`required` (Bool — обязательность согласно "Тип характеристики" или эвристике).

Парсинг операторов (≥, ≤, ±, диапазоны "X и Y") — отдельная задача `value_parser.py`
(следующая итерация). Сейчас просто сохраняем raw текст.
"""
from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

log = logging.getLogger(__name__)

# python-docx импортируется лениво — чтобы analyzer skeleton работал без него
# (для archives где DOCX вообще нет)
_DOCX_AVAILABLE: bool | None = None


def _ensure_docx() -> bool:
    global _DOCX_AVAILABLE
    if _DOCX_AVAILABLE is None:
        try:
            import docx  # noqa: F401
            _DOCX_AVAILABLE = True
        except ImportError:
            _DOCX_AVAILABLE = False
            log.warning("python-docx not installed — DOCX extractor disabled")
    return _DOCX_AVAILABLE


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

@dataclass
class ExtractedSpec:
    """Одна извлечённая характеристика."""

    name: str
    value_raw: str
    unit: str = ""
    spec_type: str = ""              # «Количественная» / «Качественная» / ''
    required: bool | None = None     # True если "Значение не может изменяться", False если "Участник указывает"
    ordinal: int | None = None       # № п/п характеристики, если есть
    notes: str = ""                  # инструкция по заполнению или дополнительная инфа

    def to_dict(self) -> dict[str, Any]:
        return {
            "name": self.name,
            "value_raw": self.value_raw,
            "unit": self.unit,
            "spec_type": self.spec_type,
            "required": self.required,
            "ordinal": self.ordinal,
            "notes": self.notes,
        }


@dataclass
class ExtractionResult:
    """Сводный результат извлечения для одного DOCX."""

    source_file: Path
    product_name: str = ""           # Из шапки или КТРУ строки
    ktru_okpd2_code: str = ""        # 32.50.50.190-00000839, если есть
    specs: list[ExtractedSpec] = field(default_factory=list)
    strategy_used: str = ""          # "ktru_table" / "kv_table" / "paragraphs" / "empty"
    notes: list[str] = field(default_factory=list)
    error: str | None = None

    def to_dict(self) -> dict[str, Any]:
        return {
            "source_file": str(self.source_file),
            "product_name": self.product_name,
            "ktru_okpd2_code": self.ktru_okpd2_code,
            "strategy_used": self.strategy_used,
            "specs": [s.to_dict() for s in self.specs],
            "notes": self.notes,
            "error": self.error,
        }


def extract_from_docx(path: Path) -> ExtractionResult:
    """Главный entrypoint — извлечь характеристики из одного DOCX."""
    result = ExtractionResult(source_file=path)

    if not _ensure_docx():
        result.error = "python-docx not installed"
        return result

    if not path.exists():
        result.error = f"file not found: {path}"
        return result

    try:
        from docx import Document
        doc = Document(str(path))
    except Exception as exc:
        result.error = f"failed to open DOCX: {exc}"
        log.warning("DOCX open failed for %s: %s", path, exc)
        return result

    # Стратегия 1: КТРУ-таблица
    for table in doc.tables:
        ktru = _try_extract_ktru_table(table)
        if ktru is not None:
            specs, product_name, okpd2 = ktru
            if specs:
                result.specs = specs
                result.product_name = product_name
                result.ktru_okpd2_code = okpd2
                result.strategy_used = "ktru_table"
                return result

    # Стратегия 2: простая key-value таблица
    for table in doc.tables:
        kv = _try_extract_kv_table(table)
        if kv:
            result.specs = kv
            result.strategy_used = "kv_table"
            # Попробуем product_name из первого параграфа
            for p in doc.paragraphs[:5]:
                if p.text.strip() and len(p.text) < 200:
                    result.product_name = p.text.strip()
                    break
            return result

    # Стратегия 3: параграфы
    para_specs = _extract_from_paragraphs(doc)
    if para_specs:
        result.specs = para_specs
        result.strategy_used = "paragraphs"
        for p in doc.paragraphs[:5]:
            if p.text.strip() and len(p.text) < 200:
                result.product_name = p.text.strip()
                break
        return result

    result.strategy_used = "empty"
    result.notes.append("no characteristics extracted from DOCX")
    return result


# ---------------------------------------------------------------------------
# Стратегия 1: КТРУ-таблица ФЗ-44
# ---------------------------------------------------------------------------

# Канонические headers КТРУ-таблиц (нормализуем для сравнения)
_KTRU_HEADERS = {
    "name": ("наименование характеристики",),
    "value": ("значение характеристики",),
    "unit": ("единица измерения характеристики", "единица измерения"),
    "type": ("тип характеристики",),
    "ordinal_char": ("№ п/п\n(хар-ки)", "№ п/п (хар-ки)", "№ п/п хар-ки", "№ п/п\nхар-ки"),
    "instruction": ("инструкция по заполнению характеристик в заявке", "инструкция по заполнению"),
    "product_name": ("наименование товара, работы, услуги", "наименование товара"),
    "ktru_code": ("ктру/окпд2", "ктру", "окпд2", "код ктру", "код окпд2"),
}

_OKPD2_RE = re.compile(r"\b\d{2}\.\d{2}\.\d{2}\.\d{3}(?:-\d{8})?\b")


def _norm_cell(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip().lower())


def _find_ktru_header_row(table: Any) -> tuple[int, dict[str, int]] | None:
    """Найти строку-заголовок и вернуть mapping column_name → col_index.

    КТРУ-таблица обычно имеет 3 строки заголовков (объединённые ячейки).
    Самая полезная — та где есть «Наименование характеристики» в чистом виде.
    """
    for r_idx, row in enumerate(table.rows[:6]):
        col_map: dict[str, int] = {}
        for c_idx, cell in enumerate(row.cells):
            normalized = _norm_cell(cell.text)
            for key, variants in _KTRU_HEADERS.items():
                if any(v in normalized for v in variants):
                    # Берём первое попадание (повторяющиеся merged cells)
                    if key not in col_map:
                        col_map[key] = c_idx
        # Если нашли минимум name + value — это header row
        if "name" in col_map and "value" in col_map:
            return r_idx, col_map
    return None


def _try_extract_ktru_table(
    table: Any,
) -> tuple[list[ExtractedSpec], str, str] | None:
    """Попытаться извлечь характеристики из таблицы как из КТРУ-формата.

    Возвращает (specs, product_name, okpd2) или None если это не КТРУ.
    """
    found = _find_ktru_header_row(table)
    if found is None:
        return None
    header_row_idx, cols = found

    # Извлекаем product_name и okpd2 из первой data row
    product_name = ""
    okpd2 = ""

    specs: list[ExtractedSpec] = []
    seen_keys: set[str] = set()

    for row in table.rows[header_row_idx + 1:]:
        # Skip rows containing только цифры порядка колонок ("1 2 3 4 ...")
        first_cells = [row.cells[i].text.strip() for i in range(min(3, len(row.cells)))]
        if all(re.fullmatch(r"\d+", c) for c in first_cells if c):
            continue

        cells = row.cells
        get = lambda key: cells[cols[key]].text.strip() if key in cols and cols[key] < len(cells) else ""

        name = get("name")
        value = get("value")
        if not name or not value:
            continue

        # Очищаем артефакты merged cells (часто дублируется в нескольких ячейках одной строки)
        if name == value:
            continue

        # Product name и OKPD2 — обычно одинаковые для всех строк характеристик одного товара
        if not product_name and "product_name" in cols:
            pn = get("product_name")
            if pn:
                product_name = pn
        if not okpd2 and "ktru_code" in cols:
            kc = get("ktru_code")
            m = _OKPD2_RE.search(kc)
            if m:
                okpd2 = m.group(0)

        unit = get("unit")
        spec_type = get("type")
        instruction = get("instruction")

        # required: если "Значение характеристики не может изменяться участником закупки"
        required: bool | None = None
        if instruction:
            instr_lower = instruction.lower()
            if "не может изменяться" in instr_lower:
                required = True
            elif "указывает" in instr_lower or "указать" in instr_lower:
                required = False

        # ordinal of characteristic
        ordinal_raw = get("ordinal_char")
        ordinal = None
        if ordinal_raw and ordinal_raw.isdigit():
            ordinal = int(ordinal_raw)

        # Защита от дублей (КТРУ-таблицы иногда повторяют строки)
        key = (ordinal, _norm_cell(name), _norm_cell(value))
        key_str = repr(key)
        if key_str in seen_keys:
            continue
        seen_keys.add(key_str)

        specs.append(ExtractedSpec(
            name=name,
            value_raw=value,
            unit=unit if unit and unit != "_" else "",
            spec_type=spec_type,
            required=required,
            ordinal=ordinal,
            notes=instruction if instruction and len(instruction) < 200 else "",
        ))

    if not specs:
        return None
    return specs, product_name, okpd2


# ---------------------------------------------------------------------------
# Стратегия 2: простая 2-3 колоночная key-value таблица
# ---------------------------------------------------------------------------

_KV_NAME_HEADERS = ("характеристика", "параметр", "наименование", "показатель")
_KV_VALUE_HEADERS = ("значение", "величина", "требование", "норма")


def _try_extract_kv_table(table: Any) -> list[ExtractedSpec]:
    """Простая key-value таблица — 2-3 колонки [имя, значение, (единица)]."""
    if not table.rows or len(table.rows) < 2:
        return []
    first_row = [_norm_cell(c.text) for c in table.rows[0].cells]
    if len(first_row) < 2 or len(first_row) > 4:
        return []

    name_col = None
    value_col = None
    unit_col = None
    for i, h in enumerate(first_row):
        if any(k in h for k in _KV_NAME_HEADERS) and name_col is None:
            name_col = i
        elif any(k in h for k in _KV_VALUE_HEADERS) and value_col is None:
            value_col = i
        elif "ед" in h and "изм" in h and unit_col is None:
            unit_col = i

    if name_col is None or value_col is None:
        # Heuristic: первые две колонки — name+value, если их ровно 2
        if len(first_row) == 2:
            name_col, value_col = 0, 1
        else:
            return []

    specs: list[ExtractedSpec] = []
    for row in table.rows[1:]:
        if len(row.cells) <= max(name_col, value_col):
            continue
        name = row.cells[name_col].text.strip()
        value = row.cells[value_col].text.strip()
        if not name or not value or name.lower() == value.lower():
            continue
        unit = row.cells[unit_col].text.strip() if unit_col is not None and unit_col < len(row.cells) else ""
        specs.append(ExtractedSpec(name=name, value_raw=value, unit=unit))
    return specs


# ---------------------------------------------------------------------------
# Стратегия 3: параграфный fallback
# ---------------------------------------------------------------------------

_PARA_KV_RE = re.compile(
    r"^([А-ЯЁA-Z][^:—–\-\n]{3,80}?)\s*[:—–\-]\s*(.+)$",
    re.MULTILINE,
)


def _extract_from_paragraphs(doc: Any) -> list[ExtractedSpec]:
    specs: list[ExtractedSpec] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text or len(text) > 500:
            continue
        m = _PARA_KV_RE.match(text)
        if not m:
            continue
        name = m.group(1).strip()
        value = m.group(2).strip()
        # Фильтр мусора — заголовки разделов и т.п.
        if len(name) < 3 or len(value) < 1:
            continue
        if name.lower().startswith(("раздел", "пункт", "статья", "приложение")):
            continue
        specs.append(ExtractedSpec(name=name, value_raw=value))
    return specs
