#!/usr/bin/env python3
"""Универсальный RAG-ingestion из MinIO -> document_registry + document_chunks + MemPalace.

Поддержка: pdf (pdfplumber), docx (python-docx), xlsx (openpyxl), rtf (striprtf), doc (antiword).
Идемпотентно: дедуп по (tenant_id, content_hash).
Тип документа классифицируется по имени файла (ТЗ/описание -> tz, заявка/предложение -> offer).

Запуск (транзиентный python:3.12-slim на сети gluvex_app_internal):
  apt-get update && apt-get install -y antiword >/dev/null;
  pip install -q pdfplumber python-docx openpyxl striprtf asyncpg minio httpx;
  python ingest_docs.py
"""
from __future__ import annotations
import os, io, re, sys, json, hashlib, asyncio, subprocess, tempfile
import asyncpg, httpx
from minio import Minio

TENANT = os.environ.get("TENANT_ID", "11111111-1111-1111-1111-111111111111")
OWNER  = os.environ.get("OWNER_ID",  "11111111-1111-1111-1111-111111111111")
EMB_MODEL = os.environ.get("EMB_MODEL", "paraphrase-multilingual-MiniLM-L12-v2")
BUCKET = os.environ.get("BUCKET", "tenders")
PREFIX = os.environ.get("PREFIX", "")
WING   = os.environ.get("WING", "gluvex-tenders")
SRC_TYPE = os.environ.get("SOURCE_TYPE", "manual_upload")
MEMPALACE = os.environ.get("MEMPALACE_URL", "http://mempalace-gluvex:8080")
CHUNK = int(os.environ.get("CHUNK_CHARS", "1200"))
OVERLAP = int(os.environ.get("CHUNK_OVERLAP", "150"))
PUSH_MP = os.environ.get("PUSH_MEMPALACE", "1") == "1"

CYR = re.compile(r"[А-Яа-яЁё]")


def classify(fn: str) -> tuple[str, str]:
    """-> (document_type, room)."""
    low = fn.lower()
    if re.search(r"предложен|заявк|форма|first|участник", low):
        return "offer", "bid"
    if re.search(r"техническое задание|\bтз\b|описание объект|приложение", low):
        return "tz", "spec"
    return "tz", "spec"  # equipment/reagent description defaults to ТЗ-like


def extract_docx(data: bytes) -> str:
    import docx
    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line.strip():
                parts.append(line)
    return "\n".join(parts)


def extract_xlsx(data: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"# Лист: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_pdf(data: bytes) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for pg in pdf.pages:
            t = pg.extract_text() or ""
            if t.strip():
                parts.append(t)
    return "\n\n".join(parts)


def extract_rtf(data: bytes) -> str:
    from striprtf.striprtf import rtf_to_text
    return rtf_to_text(data.decode("utf-8", errors="ignore"))


def extract_doc(data: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as f:
        f.write(data); path = f.name
    try:
        out = subprocess.run(["antiword", path], capture_output=True, timeout=120)
        return out.stdout.decode("utf-8", errors="ignore")
    except Exception:
        return ""
    finally:
        try: os.unlink(path)
        except Exception: pass


def extract(key: str, data: bytes) -> str:
    ext = key.rsplit(".", 1)[-1].lower()
    if ext == "pdf": return extract_pdf(data)
    if ext == "docx": return extract_docx(data)
    if ext == "xlsx": return extract_xlsx(data)
    if ext == "rtf": return extract_rtf(data)
    if ext == "doc": return extract_doc(data)
    return ""


def chunk_text(text: str, size=CHUNK, overlap=OVERLAP) -> list[str]:
    text = re.sub(r"[ \t]+", " ", text).strip()
    paras = [p.strip() for p in re.split(r"\n\s*\n|\n", text) if p.strip()]
    chunks, cur = [], ""
    for p in paras:
        if len(cur) + len(p) + 1 <= size:
            cur = (cur + "\n" + p).strip()
        else:
            if cur: chunks.append(cur)
            if len(p) <= size:
                cur = p
            else:
                for i in range(0, len(p), size - overlap):
                    chunks.append(p[i:i + size])
                cur = ""
    if cur: chunks.append(cur)
    if overlap > 0 and len(chunks) > 1:
        out = [chunks[0]]
        for i in range(1, len(chunks)):
            out.append((chunks[i - 1][-overlap:] + " " + chunks[i]).strip())
        chunks = out
    return chunks


async def main() -> int:
    mc = Minio(os.environ["MINIO_ENDPOINT"], access_key=os.environ["MINIO_ACCESS"],
               secret_key=os.environ["MINIO_SECRET"], secure=False)
    pg = await asyncpg.connect(host=os.environ.get("PGHOST", "app-db"),
                               user=os.environ.get("PGUSER", "postgres"),
                               password=os.environ["PGPASSWORD"],
                               database=os.environ.get("PGDATABASE", "gluvex_documents"))
    exts = (".pdf", ".docx", ".xlsx", ".rtf", ".doc")
    objs = [o.object_name for o in mc.list_objects(BUCKET, prefix=PREFIX, recursive=True)
            if o.object_name.lower().endswith(exts)]
    print(f"found {len(objs)} docs under {BUCKET}/{PREFIX}", flush=True)

    stats = {"docs": 0, "skipped_dup": 0, "no_text": 0, "chunks": 0, "mp_pushed": 0, "errors": 0}
    client = httpx.Client(timeout=120) if PUSH_MP else None

    for key in objs:
        try:
            data = mc.get_object(BUCKET, key).read()
            chash = hashlib.sha256(data).digest()
            if await pg.fetchval("SELECT id FROM document_registry WHERE tenant_id=$1 AND content_hash=$2", TENANT, chash):
                stats["skipped_dup"] += 1; continue
            text = extract(key, data)
            if len(text.strip()) < 30:
                stats["no_text"] += 1
                print(f"  NO TEXT: {key}", flush=True); continue
            filename = key.rsplit("/", 1)[-1]
            ext = filename.rsplit(".", 1)[-1].lower()
            ctype = {"pdf": "application/pdf", "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                     "rtf": "application/rtf", "doc": "application/msword"}.get(ext, "application/octet-stream")
            dtype, room = classify(filename)
            lang = "ru" if len(CYR.findall(text)) > 30 else "en"
            segs = key.split("/")
            person = segs[1] if len(segs) > 2 else ""
            year = next((s for s in segs if re.fullmatch(r"20\d\d.*", s)), "")
            meta = {"person": person, "year": re.sub(r"\D", "", year), "doc_kind": dtype,
                    "object_key": key, "lang": lang, "ingest": "ingest_docs_v1"}
            doc_id = await pg.fetchval("""
                INSERT INTO document_registry
                  (tenant_id, source_type, bucket, object_key, filename, content_type,
                   content_hash, document_type, language, status, access_level,
                   embedding_model, ocr_applied, owner_id, metadata, indexed_at)
                VALUES ($1,$2,$3,$4,$5,$6,$7,$8::document_type_t,$9,'actual','internal',$10,false,$11,$12, now())
                RETURNING id
            """, TENANT, SRC_TYPE, BUCKET, key, filename, ctype, chash, dtype, lang, EMB_MODEL, OWNER, json.dumps(meta))

            chunks = chunk_text(text)
            for i, ch in enumerate(chunks):
                await pg.execute("""
                    INSERT INTO document_chunks (document_id, version_id, chunk_index, chunk_text, chunk_hash, embedding_model, metadata)
                    VALUES ($1,$1,$2,$3,$4,$5,$6) ON CONFLICT (document_id, version_id, chunk_index) DO NOTHING
                """, doc_id, i, ch, hashlib.sha256(ch.encode()).digest(), EMB_MODEL,
                     json.dumps({"object_key": key, "lang": lang, "doc_kind": dtype}))
                stats["chunks"] += 1
                if PUSH_MP:
                    try:
                        r = client.post(f"{MEMPALACE}/drawer", json={
                            "content": ch, "wing": WING, "room": room,
                            "title": f"{filename} #{i}", "source_file": key,
                            "added_by": "ingest_docs", "tags": [dtype, lang, person, re.sub(r"\D","",year)]})
                        if r.status_code < 300: stats["mp_pushed"] += 1
                    except Exception as e:
                        print(f"  mp err {key}#{i}: {e}", flush=True)
            await pg.execute("UPDATE document_registry SET chunk_count=$2, updated_at=now() WHERE id=$1", doc_id, len(chunks))
            stats["docs"] += 1
            print(f"  [{stats['docs']}] {filename}  type={dtype} lang={lang} chunks={len(chunks)}", flush=True)
        except Exception as e:
            stats["errors"] += 1
            print(f"  ERROR {key}: {e}", flush=True)

    await pg.close()
    if client: client.close()
    print("=== SUMMARY ===", flush=True)
    print(json.dumps(stats, ensure_ascii=False), flush=True)
    return 0


if __name__ == "__main__":
    sys.exit(asyncio.run(main()))
