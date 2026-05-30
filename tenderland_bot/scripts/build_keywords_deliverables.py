# -*- coding: utf-8 -*-
"""
Сборка ФИНАЛЬНЫХ deliverables для менеджера, который вводит автопоиски в Tenderland руками.

Читает 4 конфига keywords_config*.md из ../config/, вытаскивает точные строки
INCLUDE/EXCLUDE из fenced-блоков (```text ... ```) и собирает:
  - Tenderland_автопоиски_ФИНАЛ.xlsx  (рабочий инструмент: копи-паст строк)
  - Tenderland_автопоиски_ФИНАЛ.docx  (объяснения + блоки для чтения)

Запуск:  python scripts/build_keywords_deliverables.py
"""
from __future__ import annotations
import re
from pathlib import Path

CONFIG = Path(__file__).resolve().parent.parent / "config"

# ---------------------------------------------------------------------------
# 1. Парсер: вытащить (heading_line, code_text) для всех ```text блоков
# ---------------------------------------------------------------------------
def get_blocks(path: Path):
    blocks = []
    heading = None
    in_fence = False
    buf = []
    for line in path.read_text(encoding="utf-8").splitlines():
        s = line.strip()
        if s.startswith("#"):
            if not in_fence:
                heading = s
        if s.startswith("```"):
            if not in_fence:
                in_fence = True
                buf = []
            else:
                in_fence = False
                blocks.append((heading, "\n".join(buf).strip()))
            continue
        if in_fence:
            buf.append(line)
    return blocks


def find_include(blocks, code):
    """INCLUDE-строка для автопоиска по его коду (ищем `code` в заголовке)."""
    needle = "`" + code + "`"
    for h, txt in blocks:
        if h and needle in h:
            return txt
    raise KeyError(f"INCLUDE not found for {code}")


def find_by_heading(blocks, *substrings):
    """Первый блок, чей заголовок содержит ВСЕ подстроки."""
    for h, txt in blocks:
        if h and all(ss in h for ss in substrings):
            return txt
    raise KeyError(f"block not found for {substrings}")


b_an = get_blocks(CONFIG / "keywords_config.md")
b_md = get_blocks(CONFIG / "keywords_config_molecular_diagnostics.md")
b_lab = get_blocks(CONFIG / "keywords_config_general_lab.md")
b_cer = get_blocks(CONFIG / "keywords_config_capillary_centrifuges_robotics.md")

# EXCLUDE-блоки (точные строки из файлов)
EXC_AN = find_by_heading(b_an, "Общая EXCLUDE")
EXC_MD_BASE = find_by_heading(b_md, "7.1")
EXC_MD_PCR = find_by_heading(b_md, "7.2")
EXC_LAB_BASE = find_by_heading(b_lab, "7.1")
EXC_LAB_DOMAIN = find_by_heading(b_lab, "7.2")
EXC_CER_BASE = find_by_heading(b_cer, "7.1")
EXC_CER_01 = find_by_heading(b_cer, "7.2")
EXC_CER_02 = find_by_heading(b_cer, "7.3")
EXC_CER_03 = find_by_heading(b_cer, "7.4")

# узкие PCR-маркеры для MDX_06 (из примечания 7.2 молекулярки)
MDX06_PCR_NARROW = "cobas Rotor-Gene LightCycler ДТпрайм CFX96 QuantStudio амплификатор"


def join(*parts):
    return "  ".join(p.strip() for p in parts if p and p.strip())


# ---------------------------------------------------------------------------
# 2. Метаданные автопоисков
# ---------------------------------------------------------------------------
# поле: code, name, group, what, brands, include, exclude_recipe, exclude_full, note, optional
SEARCHES = []


def add(code, group, what, brands, include, exclude_recipe, exclude_full,
        note="", optional=False):
    SEARCHES.append(dict(code=code, group=group, what=what, brands=brands,
                         include=include, exclude_recipe=exclude_recipe,
                         exclude_full=exclude_full, note=note, optional=optional))


GRP_AN = "Инструментальная аналитика"
GRP_MD = "Молекулярная диагностика"
GRP_LAB = "Общелабораторное"
GRP_CER = "КЭ / центрифуги / роботизация"

# ---- Аналитика ----
add("01_LC_LCMS_GPC_Prep", GRP_AN,
    "ВЭЖХ / УВЭЖХ, ЖХ-МС/МС, ГПХ (SEC/GPC), препаративная хроматография",
    "Agilent (1100–1290; 6100–6500 MS), Shimadzu (Nexera, LCMS-8060NX/9030), Waters (ACQUITY, Alliance, Xevo, SYNAPT), Thermo (Vanquish, UltiMate, Orbitrap, TSQ), PerkinElmer, SCIEX (Triple Quad, QTRAP, TripleTOF, ZenoTOF)",
    find_include(b_an, "01_LC_LCMS_GPC_Prep"), "Общий EXCLUDE (раздел 6)", EXC_AN)
add("02_GC_GCMS", GRP_AN,
    "ГХ, ГХ-МС/МС, детекторы FID/ECD/NPD/TCD, пиролизеры, парофаза/SPME",
    "Agilent 7890/8890/8860, Shimadzu GC-2030/Nexis, Thermo TRACE/ISQ/TSQ, PerkinElmer Clarus, Bruker SCION, Хроматэк-Кристалл, Хромос, Цвет",
    find_include(b_an, "02_GC_GCMS"), "Общий EXCLUDE (раздел 6)", EXC_AN)
add("03_ICP_OES", GRP_AN, "ИСП-ОЭС / ИСП-АЭС (оптико-эмиссионная с плазмой)",
    "Agilent 5800/5900/5100/5110, PerkinElmer Avio/Optima, Thermo iCAP, Spectro Arcos/Genesis, Shimadzu ICPE-9800, Analytik Jena PlasmaQuant",
    find_include(b_an, "03_ICP_OES"), "Общий EXCLUDE (раздел 6)", EXC_AN)
add("04_AAS", GRP_AN, "Атомно-абсорбционная спектрометрия (пламя + ЭТА/графит)",
    "Agilent 240/280/AA240Z, PerkinElmer PinAAcle/AAnalyst, Shimadzu AA-7000/7800, Thermo iCE 3000, Analytik Jena contrAA/ZEEnit/novAA, МГА-915/1000, КВАНТ",
    find_include(b_an, "04_AAS"), "Общий EXCLUDE (раздел 6)", EXC_AN)
add("05_ICP_MS", GRP_AN, "ИСП-МС, тандем QQQ, секторные HR-ICP-MS",
    "Agilent 7700–8900, PerkinElmer NexION, Thermo iCAP Q/RQ/TQ, Element 2/XR, Neptune Plus, Shimadzu ICPMS-2030",
    find_include(b_an, "05_ICP_MS"), "Общий EXCLUDE (раздел 6)", EXC_AN)
add("06_IC", GRP_AN, "Ионная хроматография (анионная/катионная, кондуктометрия, супрессор)",
    "Thermo Dionex ICS-1100..6000, Metrohm 882/930/940",
    find_include(b_an, "06_IC"), "Общий EXCLUDE (раздел 6)", EXC_AN)
add("07_UV_Vis", GRP_AN, "УФ-видимая спектрофотометрия (двухлучевые, диод-матрица, NIR)",
    "Agilent Cary 60–7000, Shimadzu UV-1900/2700/3600, Thermo Evolution/NanoDrop, PerkinElmer Lambda",
    find_include(b_an, "07_UV_Vis"), "Общий EXCLUDE (раздел 6)", EXC_AN)
add("08_FTIR", GRP_AN, "ФТИР / ИК-Фурье, ИК-микроскопия, НПВО/ATR",
    "Agilent Cary 630–680, Shimadzu IRTracer/IRAffinity, Thermo Nicolet, PerkinElmer Spectrum/Frontier",
    find_include(b_an, "08_FTIR"), "Общий EXCLUDE (раздел 6)", EXC_AN)
add("09_Service", GRP_AN, "Сервис / поверка / IQ-OQ-PQ на приборы 01–08",
    "Любые из 01–08",
    find_include(b_an, "09_Service"),
    "СМЯГЧЁННЫЙ EXCLUDE: из общего убрать «реактив», «реагент», «гель» (в сервисе часто всплывают)",
    EXC_AN, optional=True)
add("10_Consumables", GRP_AN, "Расходники: колонки, картриджи, лампы, кюветы, ГСО",
    "Любые из 01–08 + Dionex, Metrohm, Knauer, JASCO",
    find_include(b_an, "10_Consumables"),
    "СМЯГЧЁННЫЙ EXCLUDE: снять «реактив», «реагент», «гель» (расходники часто их содержат)",
    EXC_AN, optional=True)

# ---- Молекулярка ----
EXC_MD = join(EXC_MD_BASE, EXC_MD_PCR)
add("MDX_01_Sequencers", GRP_MD,
    "NGS-секвенаторы + Sanger CE / капиллярные генетические анализаторы (только ПРИБОРЫ)",
    "Illumina (iSeq..NovaSeq X+), MGI DNBSEQ (G50..T20), GeneMind, Salus, Хеликон, Р-Ген, ONT (MinION/PromethION), PacBio (Revio), Element AVITI, ABI 3500/3500xL/3130/3730, SeqStudio",
    find_include(b_md, "MDX_01_Sequencers"),
    "Базовый (7.1) + ПЦР/ИФА/single-cell (7.2)", EXC_MD)
add("MDX_02_Reagents_Libraries", GRP_MD,
    "Расходники секвенирования, library prep, выделение НК, авто-станции",
    "Illumina TruSeq/Nextera, MGIEasy, KAPA, NEBNext, QIAseq, Twist, IDT xGen, Agilent SureSelect, KingFisher, MGISP, Vazyme, Pillar, Burning Rock, Novogene, BigDye/POP-7 (Sanger)",
    find_include(b_md, "MDX_02_Reagents_Libraries"),
    "Базовый (7.1) + ПЦР/ИФА/single-cell (7.2)", EXC_MD)
add("MDX_03_Oncology_Panels", GRP_MD,
    "Онкопанели и тесты: BRCA/EGFR/KRAS/BRAF/ALK/MSI/HRD/TMB/CGP",
    "AmoyDx, Pillar, Burning Rock, Novogene, Parseq, OncoAtlas, TestGen, FoundationOne, Oncomine, Archer, TSO500, Twist, Personalis, Tempus, Caris",
    find_include(b_md, "MDX_03_Oncology_Panels"),
    "Базовый (7.1) + ПЦР/ИФА/single-cell (7.2)", EXC_MD)
add("MDX_04_NIPT_PGT_HLA", GRP_MD,
    "НИПТ, ПГТ-А/М, HLA-типирование, биочипы",
    "Illumina VeriSeq, Vanadis, BambniTest, LinkSeq/AllType/Olerup (HLA), Affymetrix CytoScan, Infinium BeadChip",
    find_include(b_md, "MDX_04_NIPT_PGT_HLA"),
    "Базовый (7.1) + ПЦР/ИФА/single-cell (7.2)", EXC_MD)
add("MDX_06_Sequencing_Services", GRP_MD,
    "УСЛУГИ секвенирования: CES/WES/WGS, биоинф. интерпретация, аутсорс NGS",
    "Genotek, Атлас Биомед, Парсек-Лаб, Хеликон-CRO, Macrogen, BGI-CRO, Novogene-services, Foundation Medicine как услуга",
    find_include(b_md, "MDX_06_Sequencing_Services"),
    "Базовый (7.1). ПЦР-блок (7.2) НЕ ставить целиком — в услугах ПЦР упоминается как этап pipeline. "
    "Если шумно — добавить только узкие маркеры: " + MDX06_PCR_NARROW,
    EXC_MD_BASE,
    note="Новый кластер (рекомендован к включению). Тип лота — услуги (ОКПД2 71.20 / 86.10).")
add("MDX_05_Service", GRP_MD,
    "Сервис ТОЛЬКО Illumina / MGI-DNBSEQ / ABI 3500 (узкий скоуп)",
    "Illumina, MGI/DNBSEQ, ABI 3500/3500xL, SeqStudio",
    find_include(b_md, "MDX_05_Service"),
    "Только базовый (7.1). ПЦР-блок (7.2) НЕ нужен — скоуп и так ограничен моделями.",
    EXC_MD_BASE, optional=True)

# ---- Общелаб ----
EXC_LAB = join(EXC_LAB_BASE, EXC_LAB_DOMAIN)
add("LAB_01_Climate", GRP_LAB,
    "CO2-инкубаторы, термостаты, сушильные шкафы, климатокамеры, муфельные/трубчатые печи",
    "Memmert, Binder, Thermo Heracell/Heratherm, Lauda, Julabo, Huber, Nabertherm, Carbolite, СНОЛ, ШС, Eppendorf ThermoMixer",
    find_include(b_lab, "LAB_01_Climate"),
    "Базовый (7.1) + отсечь аналитику/молекулярку (7.2)", EXC_LAB)
add("LAB_02_Sterilization", GRP_LAB,
    "Паровые автоклавы, сухожаровые / плазменные / газовые стерилизаторы",
    "Tuttnauer, Systec, Steris, Memmert SF, ВК-30/75, ГК-100, ГП-560, ШСС, Sterrad",
    find_include(b_lab, "LAB_02_Sterilization"),
    "Базовый (7.1) + отсечь аналитику/молекулярку (7.2)", EXC_LAB)
add("LAB_03_Evaporation", GRP_LAB,
    "Роторные испарители, концентраторы, лиофильные сушилки, ультразвук",
    "Buchi Rotavapor, Heidolph Hei-VAP, IKA RV, Christ Alpha, Labconco FreeZone, Eppendorf Concentrator, Elma, Bandelin, Hielscher",
    find_include(b_lab, "LAB_03_Evaporation"),
    "Базовый (7.1) + отсечь аналитику/молекулярку (7.2)", EXC_LAB)
add("LAB_04_Mixing_Homogenization", GRP_LAB,
    "Мешалки (магн./верхнепривод), шейкеры, вортексы, гомогенизаторы, мельницы",
    "IKA, Heidolph, Velp, Stuart, Eppendorf, Retsch, Fritsch, BioSan, Корвет",
    find_include(b_lab, "LAB_04_Mixing_Homogenization"),
    "Базовый (7.1) + отсечь аналитику/молекулярку (7.2)", EXC_LAB)
add("LAB_05_Reactors", GRP_LAB,
    "Биореакторы/ферментёры, химические реакторы, синтезаторы, микроволновые",
    "Sartorius BIOSTAT/ambr, Eppendorf BioFlo/DASGIP, INFORS, Applikon, Cytiva, Buchi, Asynt, Radleys, Mettler OptiMax/EasyMax, Anton Paar, CEM, Milestone, Parr, Berghof",
    find_include(b_lab, "LAB_05_Reactors"),
    "Базовый (7.1) + отсечь аналитику/молекулярку (7.2)", EXC_LAB)
add("LAB_06_Weighing_Water_pH", GRP_LAB,
    "Аналитические/прецизионные весы, системы очистки воды (Milli-Q), pH/EC/O2-метры",
    "Mettler-Toledo, Sartorius, OHAUS, Kern, Millipore Milli-Q, ELGA, Hanna, WTW",
    find_include(b_lab, "LAB_06_Weighing_Water_pH"),
    "Базовый (7.1) + отсечь аналитику/молекулярку (7.2). Доп. отсечь торговые/бытовые весы.",
    EXC_LAB, optional=True)

# ---- КЭ / центрифуги / роботизация ----
add("CER_01_Capillary_Electrophoresis", GRP_CER,
    "Капиллярный электрофорез НЕ для Sanger (CZE/MEKC/CIEF/CGE) — фарма/биохимия",
    "Beckman PA 800 Plus / P-ACE / CESI 8000, Sciex BioPhase 8800, Agilent 7100 CE, Bio-Rad BioFocus, Lumex Капель",
    find_include(b_cer, "CER_01_Capillary_Electrophoresis"),
    "Базовый (7.1) + отсечь Sanger CE и плоский гель-электрофорез (7.2)",
    join(EXC_CER_BASE, EXC_CER_01))
add("CER_02_Centrifuges", GRP_CER,
    "Лабораторные центрифуги всех классов (настольные, рефрижераторные, ультра, минифуги)",
    "Eppendorf 5418–5920R, Hettich, Sigma, Thermo Sorvall/Heraeus, Beckman Avanti/Optima/Allegra, Hermle, Kubota, Hitachi, Liston, ELMI, BioSan, ОПН",
    find_include(b_cer, "CER_02_Centrifuges"),
    "Базовый (7.1) + отсечь промышленные и медицинские-для-крови (7.3)",
    join(EXC_CER_BASE, EXC_CER_02))
add("CER_03_Liquid_Handling_Robotics", GRP_CER,
    "Liquid handling роботы, dispensers, автоматизированная NGS-пробоподготовка",
    "Tecan EVO/Fluent, Hamilton STAR/Vantage/NIMBUS, Beckman Biomek/Echo, Eppendorf epMotion, Opentrons, Agilent Bravo/AssayMAP, PerkinElmer JANUS/Sciclone, MGI MGISP, Vazyme, Nanodigm IsoFlux",
    find_include(b_cer, "CER_03_Liquid_Handling_Robotics"),
    "Базовый (7.1) + отсечь индустриальную робототехнику KUKA/ABB/Fanuc (7.4)",
    join(EXC_CER_BASE, EXC_CER_03))

# Глобальные параметры
GLOBAL_PARAMS = [
    ("Фильтр «Включать»", "tender_keywords_include (id 136, тип text)", "Полнотекст по name + notification + files"),
    ("Фильтр «Исключать»", "tender_keywords_exclude (id 137, тип text)", "То же поле"),
    ("Тип закупки", "Все: 44-ФЗ, 223-ФЗ, коммерческие, СНГ", "Не ограничивать"),
    ("Регион", "Все регионы РФ + СНГ", "Не ограничивать"),
    ("Дата публикации", "range: от = СЕГОДНЯ−7, до = пусто", "Скользящее окно с буфером на сбои"),
    ("Дата окончания подачи", "range: от = СЕГОДНЯ, до = пусто", "Только активные тендеры"),
    ("Сортировка", "tender_sysPublishDate.desc", "Новые сверху"),
    ("Ключ дедупликации (CLI)", "tender_id (формат TL*)", "НЕ regNumber"),
    ("Частота прогона", "1 раз в сутки в рабочие дни (cron 0 8 * * 1-5)", "Внутри free-tier API"),
]

SYNTAX = [
    ("(пробел)", "OR — любой из терминов", "ВЭЖХ HPLC  → один из двух"),
    ("++", "AND со стеммингом (оба слова в любых формах)", "жидкостн++хроматограф"),
    ("+", "соединение в точную фразу", "тройн+квадрупол"),
    ("=", "точное совпадение, отключает русскую морфологию", "=HPLC  =ICP-MS  =PGT-A"),
]

STEPS = [
    "Создать новый автопоиск с именем из колонки «Имя автопоиска» (например 01_LC_LCMS_GPC_Prep).",
    "Добавить ОДИН фильтр «Поиск по ключевым словам». В «Включать» вставить строку INCLUDE, в «Исключать» — строку EXCLUDE (целиком, одной строкой).",
    "Добавить фильтр «Дата публикации» (range): от = СЕГОДНЯ−7 дней, до = пусто.",
    "Добавить фильтр «Дата окончания подачи» (range): от = СЕГОДНЯ, до = пусто.",
    "Сохранить → скопировать id из URL (?id=XXXXX).",
    "Прислать список пар «имя → id» (для внесения в autosearches.toml).",
]

WARNINGS = [
    "Каждую строку вставлять ОДНОЙ строкой, без переносов. После вставки проверить, что внутрь не попал \\n.",
    "Tenderland хранит АБСОЛЮТНЫЕ даты — раз в неделю вручную обновлять «от» в фильтрах дат.",
    "Дубли одного тендера между автопоисками — это нормально, CLI дедуплицирует по tender_id.",
    "Через 1 неделю работы пройтись по «Чеклисту после первого прогона» в исходных .md и подкрутить EXCLUDE.",
]


# ===========================================================================
# 3. EXCEL
# ===========================================================================
def build_xlsx():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()

    NAVY = "1F3864"; BLUE = "2E5496"; LIGHT = "D9E1F2"; GREEN = "E2EFDA"
    GREY = "F2F2F2"; AMBER = "FFF2CC"
    hdr = Font(bold=True, color="FFFFFF", size=11)
    title = Font(bold=True, color="FFFFFF", size=14)
    bold = Font(bold=True)
    mono = Font(name="Consolas", size=9)
    thin = Side(style="thin", color="BFBFBF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    wrap = Alignment(wrap_text=True, vertical="top")
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)

    def fill(hexc):
        return PatternFill("solid", fgColor=hexc)

    # ---- Лист 1: Инструкция ----
    ws = wb.active
    ws.title = "Инструкция"
    ws.sheet_view.showGridLines = False
    ws.column_dimensions["A"].width = 3
    ws.column_dimensions["B"].width = 30
    ws.column_dimensions["C"].width = 70
    ws.column_dimensions["D"].width = 45

    r = 1
    ws.merge_cells(f"B{r}:D{r}")
    c = ws[f"B{r}"]; c.value = "Tenderland — автопоиски Глювекса. Инструкция для менеджера"
    c.font = title; c.fill = fill(NAVY); c.alignment = center
    ws.row_dimensions[r].height = 32
    r += 2

    def section(text):
        nonlocal r
        ws.merge_cells(f"B{r}:D{r}")
        cc = ws[f"B{r}"]; cc.value = text; cc.font = hdr; cc.fill = fill(BLUE)
        cc.alignment = Alignment(vertical="center")
        ws.row_dimensions[r].height = 22
        r += 1

    def kv_table(rows, c1="Параметр", c2="Значение", c3="Комментарий"):
        nonlocal r
        for i, name in enumerate((c1, c2, c3)):
            cc = ws.cell(row=r, column=2 + i, value=name)
            cc.font = bold; cc.fill = fill(LIGHT); cc.border = border; cc.alignment = wrap
        r += 1
        for row in rows:
            for i, val in enumerate(row):
                cc = ws.cell(row=r, column=2 + i, value=val)
                cc.border = border; cc.alignment = wrap
            r += 1
        r += 1

    section("Что это и как пользоваться")
    intro = [
        ("Файл", "Назначение", ""),
        ("Этот Excel", "Рабочий инструмент: на листе «INCLUDE+EXCLUDE» — готовые строки для копи-паста в поля «Включать» / «Исключать».", ""),
        ("Word-версия", "То же самое + развёрнутые объяснения по каждому блоку (для чтения и онбординга).", ""),
    ]
    for row in intro:
        for i, val in enumerate(row):
            cc = ws.cell(row=r, column=2 + i, value=val)
            cc.border = border; cc.alignment = wrap
            if r == 4 + 1:
                cc.font = bold
        r += 1
    r += 1
    # fix header styling for intro
    for i in range(2):
        ws.cell(row=4, column=2 + i).font = bold
        ws.cell(row=4, column=2 + i).fill = fill(LIGHT)

    section("Синтаксис ключевых слов Tenderland")
    kv_table(SYNTAX, "Оператор", "Значение", "Пример")

    section("Порядок настройки одного автопоиска (повторить для каждого)")
    for i, step in enumerate(STEPS, 1):
        ws.cell(row=r, column=2, value=f"Шаг {i}").font = bold
        ws.cell(row=r, column=2).border = border; ws.cell(row=r, column=2).alignment = wrap
        ws.merge_cells(f"C{r}:D{r}")
        cc = ws.cell(row=r, column=3, value=step); cc.border = border; cc.alignment = wrap
        ws.cell(row=r, column=4).border = border
        r += 1
    r += 1

    section("Глобальные параметры (одинаковы для всех автопоисков)")
    kv_table(GLOBAL_PARAMS)

    section("Важно / частые ошибки")
    for w in WARNINGS:
        ws.merge_cells(f"B{r}:D{r}")
        cc = ws.cell(row=r, column=2, value="•  " + w)
        cc.fill = fill(AMBER); cc.border = border; cc.alignment = wrap
        ws.row_dimensions[r].height = 30
        r += 1

    # ---- Лист 2: Карта автопоисков ----
    ws2 = wb.create_sheet("Карта автопоисков")
    ws2.sheet_view.showGridLines = False
    cols2 = [("№", 5), ("Имя автопоиска", 34), ("Группа", 28), ("Что ищем", 60),
             ("Статус", 14)]
    for i, (name, w) in enumerate(cols2, 1):
        ws2.column_dimensions[get_column_letter(i)].width = w
    ws2.merge_cells("A1:E1")
    t = ws2["A1"]; t.value = "Карта всех автопоисков (19 основных + 4 опциональных)"
    t.font = title; t.fill = fill(NAVY); t.alignment = center
    ws2.row_dimensions[1].height = 28
    hr = 2
    for i, (name, _) in enumerate(cols2, 1):
        cc = ws2.cell(row=hr, column=i, value=name)
        cc.font = hdr; cc.fill = fill(BLUE); cc.border = border; cc.alignment = center
    rr = hr + 1
    for idx, s in enumerate(SEARCHES, 1):
        vals = [idx, s["code"], s["group"], s["what"],
                "опц." if s["optional"] else "основной"]
        for i, v in enumerate(vals, 1):
            cc = ws2.cell(row=rr, column=i, value=v)
            cc.border = border; cc.alignment = wrap if i in (2, 3, 4) else center
            cc.fill = fill(GREY if s["optional"] else "FFFFFF")
        rr += 1
    ws2.freeze_panes = "A3"

    # ---- Лист 3: INCLUDE + EXCLUDE (главный рабочий) ----
    ws3 = wb.create_sheet("INCLUDE+EXCLUDE")
    ws3.sheet_view.showGridLines = False
    cols3 = [("№", 4), ("Имя автопоиска", 30), ("Что ищем / бренды", 40),
             ("INCLUDE — вставить в «Включать»", 90),
             ("Какой EXCLUDE", 34),
             ("EXCLUDE — вставить в «Исключать»", 80)]
    for i, (name, w) in enumerate(cols3, 1):
        ws3.column_dimensions[get_column_letter(i)].width = w
    ws3.merge_cells("A1:F1")
    t = ws3["A1"]; t.value = "Готовые строки для копи-паста (INCLUDE → «Включать», EXCLUDE → «Исключать»)"
    t.font = title; t.fill = fill(NAVY); t.alignment = center
    ws3.row_dimensions[1].height = 28
    hr = 2
    for i, (name, _) in enumerate(cols3, 1):
        cc = ws3.cell(row=hr, column=i, value=name)
        cc.font = hdr; cc.fill = fill(BLUE); cc.border = border; cc.alignment = center
    rr = hr + 1
    cur_group = None
    for idx, s in enumerate(SEARCHES, 1):
        if s["group"] != cur_group:
            cur_group = s["group"]
            ws3.merge_cells(start_row=rr, start_column=1, end_row=rr, end_column=6)
            gc = ws3.cell(row=rr, column=1, value="▎ " + cur_group)
            gc.font = Font(bold=True, color=NAVY, size=11); gc.fill = fill(LIGHT)
            gc.alignment = Alignment(vertical="center")
            rr += 1
        whatbr = s["what"] + "\n\nБренды: " + s["brands"]
        if s["note"]:
            whatbr += "\n\n⚑ " + s["note"]
        vals = [idx, s["code"] + ("  (опц.)" if s["optional"] else ""),
                whatbr, s["include"], s["exclude_recipe"], s["exclude_full"]]
        for i, v in enumerate(vals, 1):
            cc = ws3.cell(row=rr, column=i, value=v)
            cc.border = border
            if i in (4, 6):
                cc.font = mono; cc.alignment = wrap
            elif i == 1:
                cc.alignment = center
            else:
                cc.alignment = wrap
            if s["optional"]:
                cc.fill = fill(GREY)
        ws3.row_dimensions[rr].height = 150
        rr += 1
    ws3.freeze_panes = "A3"

    # ---- Лист 4: EXCLUDE-блоки (справочник) ----
    ws4 = wb.create_sheet("EXCLUDE-блоки")
    ws4.sheet_view.showGridLines = False
    ws4.column_dimensions["A"].width = 38
    ws4.column_dimensions["B"].width = 55
    ws4.column_dimensions["C"].width = 95
    ws4.merge_cells("A1:C1")
    t = ws4["A1"]; t.value = "Справочник EXCLUDE-блоков (из них собраны строки на листе INCLUDE+EXCLUDE)"
    t.font = title; t.fill = fill(NAVY); t.alignment = center
    ws4.row_dimensions[1].height = 28
    for i, name in enumerate(("Блок", "Где применяется", "Строка"), 1):
        cc = ws4.cell(row=2, column=i, value=name)
        cc.font = hdr; cc.fill = fill(BLUE); cc.border = border; cc.alignment = center
    exc_rows = [
        ("Аналитика — общий (разд. 6)", "Все 01–08 (для 09/10 — смягчить)", EXC_AN),
        ("Молекулярка — базовый (7.1)", "Все MDX_01–06", EXC_MD_BASE),
        ("Молекулярка — ПЦР/ИФА/single-cell (7.2)", "MDX_01–04 (НЕ MDX_05/06 целиком)", EXC_MD_PCR),
        ("Молекулярка — узкие PCR-маркеры", "MDX_06, если шумно", MDX06_PCR_NARROW),
        ("Общелаб — базовый (7.1)", "Все LAB_01–06", EXC_LAB_BASE),
        ("Общелаб — отсечь аналитику/молекулярку (7.2)", "Все LAB_01–06", EXC_LAB_DOMAIN),
        ("КЭ/центрифуги/роботы — базовый (7.1)", "Все CER_01–03", EXC_CER_BASE),
        ("CER_01 — Sanger + плоский гель (7.2)", "CER_01", EXC_CER_01),
        ("CER_02 — промышл./мед.-кровь (7.3)", "CER_02", EXC_CER_02),
        ("CER_03 — индустр. робототехника (7.4)", "CER_03", EXC_CER_03),
    ]
    rr = 3
    for name, where, txt in exc_rows:
        a = ws4.cell(row=rr, column=1, value=name); a.font = bold
        b = ws4.cell(row=rr, column=2, value=where)
        cc = ws4.cell(row=rr, column=3, value=txt); cc.font = mono
        for col in (1, 2, 3):
            ws4.cell(row=rr, column=col).border = border
            ws4.cell(row=rr, column=col).alignment = wrap
        ws4.row_dimensions[rr].height = 90
        rr += 1
    ws4.freeze_panes = "A3"

    out = CONFIG / "Tenderland_autosearches_FINAL.xlsx"
    wb.save(out)
    return out


# ===========================================================================
# 4. WORD
# ===========================================================================
def build_docx():
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY = RGBColor(0x1F, 0x38, 0x64)
    BLUE = RGBColor(0x2E, 0x54, 0x96)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    def shade(cell, hexc):
        tcPr = cell._tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
        sh.set(qn("w:fill"), hexc); tcPr.append(sh)

    def code_block(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.2)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = "Consolas"; run.font.size = Pt(8)
        # лёгкая заливка через таблицу-обёртку было бы сложнее; оставим моноширинным
        return p

    def codebox(label, text, fillhex):
        """Однострочный код в одноячеечной таблице с заливкой."""
        t = doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = t.cell(0, 0)
        shade(cell, fillhex)
        cell.width = Cm(17)
        p = cell.paragraphs[0]
        if label:
            lr = p.add_run(label + "\n"); lr.bold = True; lr.font.size = Pt(9)
        run = p.add_run(text)
        run.font.name = "Consolas"; run.font.size = Pt(7.5)
        # тонкая граница
        tblPr = t._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
            e.set(qn("w:color"), "BFBFBF")
            borders.append(e)
        tblPr.append(borders)
        doc.add_paragraph()

    # Титул
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run("Tenderland — автопоиски Глювекса")
    run.bold = True; run.font.size = Pt(22); run.font.color.rgb = NAVY
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Финальная инструкция для менеджера: блоки ключевых слов и правила ручного ввода")
    sr.font.size = Pt(12); sr.font.color.rgb = BLUE
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("19 основных + 4 опциональных автопоиска · 4 группы оборудования\n"
                 "Источник: keywords_config*.md (v1.1, 2026-05-14)").font.size = Pt(9)
    doc.add_paragraph()

    # Как пользоваться
    doc.add_heading("Как пользоваться", level=1)
    doc.add_paragraph(
        "Для каждого автопоиска ниже даны две строки: INCLUDE (вставляется в поле «Включать») "
        "и EXCLUDE (в поле «Исключать»). Обе строки уже собраны целиком — копируйте «как есть», "
        "одной строкой, без переносов. Excel-версия (Tenderland_автопоиски_ФИНАЛ.xlsx) удобнее "
        "для самого копи-паста; этот документ — для понимания, что и зачем.")

    # Синтаксис
    doc.add_heading("Синтаксис ключевых слов", level=1)
    tbl = doc.add_table(rows=1, cols=3); tbl.style = "Light Grid Accent 1"
    for i, txt in enumerate(("Оператор", "Значение", "Пример")):
        c = tbl.rows[0].cells[i]; c.paragraphs[0].add_run(txt).bold = True
    for op, mean, ex in SYNTAX:
        cells = tbl.add_row().cells
        cells[0].paragraphs[0].add_run(op).font.name = "Consolas"
        cells[1].text = mean
        cells[2].paragraphs[0].add_run(ex).font.name = "Consolas"

    # Шаги
    doc.add_heading("Порядок настройки одного автопоиска", level=1)
    for i, step in enumerate(STEPS, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)

    # Глобальные параметры
    doc.add_heading("Глобальные параметры (для всех автопоисков)", level=1)
    tbl = doc.add_table(rows=1, cols=3); tbl.style = "Light Grid Accent 1"
    for i, txt in enumerate(("Параметр", "Значение", "Комментарий")):
        tbl.rows[0].cells[i].paragraphs[0].add_run(txt).bold = True
    for a, bv, cv in GLOBAL_PARAMS:
        cells = tbl.add_row().cells
        cells[0].text = a; cells[1].text = bv; cells[2].text = cv

    # Важно
    doc.add_heading("Важно / частые ошибки", level=1)
    for w in WARNINGS:
        doc.add_paragraph(w, style="List Bullet")

    doc.add_page_break()

    # По группам и автопоискам
    cur_group = None
    for idx, s in enumerate(SEARCHES, 1):
        if s["group"] != cur_group:
            cur_group = s["group"]
            doc.add_heading(cur_group, level=1)
        title = f"{s['code']}" + ("  (опционально)" if s["optional"] else "")
        doc.add_heading(title, level=2)
        p = doc.add_paragraph(); p.add_run("Что ищем: ").bold = True; p.add_run(s["what"])
        p = doc.add_paragraph(); p.add_run("Ключевые бренды: ").bold = True; p.add_run(s["brands"])
        if s["note"]:
            p = doc.add_paragraph(); rr = p.add_run("⚑ " + s["note"]); rr.italic = True
        codebox("INCLUDE → поле «Включать»:", s["include"], "E2EFDA")
        p = doc.add_paragraph(); p.add_run("EXCLUDE: ").bold = True; p.add_run(s["exclude_recipe"])
        codebox("EXCLUDE → поле «Исключать»:", s["exclude_full"], "FCE4D6")

    out = CONFIG / "Tenderland_autosearches_FINAL.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    x = build_xlsx()
    print("xlsx ->", x)
    d = build_docx()
    print("docx ->", d)
    print("searches:", len(SEARCHES))
