# -*- coding: utf-8 -*-
"""Сборка Word-файла v2: пилотная перенастройка автопоисков 01_LC и MDX_01.

Берёт текущие INCLUDE/EXCLUDE из _api_check/<id>_include|exclude.txt
(сняты с UI-версии пользователя) и пристраивает доп-EXCLUDE по наблюдённому шуму,
плюс инструкции для UI и проект Tier-2 LLM.

Запуск:
    python scripts/build_v2_pilots_docx.py
Выход:
    config/Tenderland_v2_pilots_01LC_MDX01.docx
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
AC = ROOT / "_api_check"
OUT = ROOT / "config" / "Tenderland_v2_pilots_01LC_MDX01.docx"

# Текущие строки (как у пользователя в UI)
INC_01 = (AC / "373343_include.txt").read_text(encoding="utf-8").strip()
EXC_01 = (AC / "373343_exclude.txt").read_text(encoding="utf-8").strip()
INC_MDX = (AC / "373417_include.txt").read_text(encoding="utf-8").strip()
EXC_MDX = (AC / "373417_exclude.txt").read_text(encoding="utf-8").strip()

# Дополнения EXCLUDE на основе шума из топ-20 свежих (см. _api_check/sample_373343.md и sample_373417.md)
EXC_01_ADD = (
    "изоляционн пирометр розетк гигрометр бытов++термометр лабораторн++термометр "
    "водонагреват двер++металл термошкаф Тахион интерфейс++защит "
    "термографическ термобель трекинг ботинк КИПиА сувенирн термоформов "
    "горелк++крематор газоочистк термическ++анализ Netzsch термоустойчив "
    "термоэлектрич термоплёнк термоплен бан++водян пиролитическ++ТМЦ "
    "ОМ++МК заглушк термопар термос++бытов"
)
EXC_MDX_ADD = (
    "дорожн++знак ядерн ДПО повышен++квалификац детск++площадк ремонт++кровл "
    "LECO++CS сер++углерод++LECO коленчат++вал ремень++генератор видеостен "
    "экскурсионн гирлянд иллюминац дизельн++топлив электромонтажн "
    "биохимическ++анализатор автоматическ++биохим IT++закупк новогодн "
    "праздничн++иллюминац шкив++коленчат"
)

EXC_01_NEW = EXC_01 + " " + EXC_01_ADD
EXC_MDX_NEW = EXC_MDX + " " + EXC_MDX_ADD


# ---------- doc helpers ----------
def shade(cell, hexc: str):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexc)
    tcPr.append(sh)


def thin_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "BFBFBF")
        borders.append(e)
    tblPr.append(borders)


def codebox(doc, label: str, text: str, fillhex: str):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shade(cell, fillhex)
    cell.width = Cm(17)
    p = cell.paragraphs[0]
    if label:
        lr = p.add_run(label + "\n")
        lr.bold = True
        lr.font.size = Pt(9)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(7.5)
    thin_borders(t)
    doc.add_paragraph()


def kv_table(doc, header, rows):
    tbl = doc.add_table(rows=1, cols=len(header))
    tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for r in rows:
        cells = tbl.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = str(v)


# ---------- build ----------
NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x54, 0x96)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

# Title
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = h.add_run("Tenderland — пилотная перенастройка автопоисков (v2)")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("01_LC_LCMS_GPC_Prep и MDX_01_Sequencers — чистка шума + Tier-2 LLM")
sr.font.size = Pt(12)
sr.font.color.rgb = BLUE

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(
    "Источник INCLUDE/EXCLUDE — твои текущие версии (id 373343, 373417), снятые через API.\n"
    "Дополнения EXCLUDE — на основе наблюдённого шума в топ-20 свежих тендеров каждого автопоиска."
).font.size = Pt(9)
doc.add_paragraph()

# 1. Контекст
doc.add_heading("1. Что увидели", level=1)
doc.add_paragraph(
    "Проверили top-20 свежайших тендеров (sort: дата публикации desc) для каждого автопоиска через "
    "API Export. Результат:"
)
kv_table(doc, ["Автопоиск", "total_count", "Релевант в top-20", "Доля шума"], [
    ["01_LC_LCMS_GPC_Prep (373343)", "985 (на потолок ~1000)", "0 из 20", "100%"],
    ["MDX_01_Sequencers (373417)", "176", "~7 из 20", "~65%"],
])
doc.add_paragraph(
    "Корень шума: автопоиск ищет одновременно «по Названию тендера» И «по документации» "
    "(tender_files). Любое слово вроде «термошкаф», «термопара», «термопрофиль» в спецификации "
    "к закупке трекинговых ботинок цепляет автопоиск 01_LC. По названию такие тендеры не цепляются."
)

# 2. Что меняем — общие принципы
doc.add_heading("2. Что меняем — общие принципы", level=1)
doc.add_paragraph(
    "Применяется к обоим пилотным автопоискам (потом раскатим на все 21):"
)
ul = doc.add_paragraph(
    "1) Поиск ТОЛЬКО по «Название тендера» (tender_name). "
    "«По документации» (tender_files) — отключить.\n"
    "Гипотеза: даёт −90% шума и не теряет существенно по релевантности, потому что серьёзные "
    "лоты на приборы упоминают тип прибора в названии (хроматограф/ВЭЖХ/HPLC/секвенатор/NGS)."
)
doc.add_paragraph(
    "2) Расширенный EXCLUDE по конкретным шумам, наблюдённым в выдаче (см. ниже строки)."
)
doc.add_paragraph(
    "3) Дата публикации: range «от = ВЧЕРА 00:00» — «до = пусто». Обновлять придётся вручную в UI; "
    "до настройки автоматизации это даёт 1-дневное окно и убирает 1000-потолок."
)
doc.add_paragraph(
    "4) НМЦК (начальная цена): «от = 4 000 000 ₽». Применять только к ПРИБОРНЫМ автопоискам "
    "(01–08, MDX_01, LAB_01–05, CER_01–02). На расходники/сервис (09, 10, MDX_02–06, LAB_06) "
    "порог НЕ ставить или ставить меньший (например ≥ 500 000 ₽)."
)
doc.add_paragraph(
    "5) Статус: «активные тендеры» (если есть такая кнопка) — уже стоит."
)

# 3. 01_LC_LCMS_GPC_Prep
doc.add_heading("3. Автопоиск 01_LC_LCMS_GPC_Prep (id 373343)", level=1)
doc.add_heading("3.1. INCLUDE (без изменений — твоя версия)", level=2)
codebox(doc, f"INCLUDE — поле «По названию тендера». Длина: {len(INC_01)} знаков.",
        INC_01, "E2EFDA")

doc.add_heading("3.2. EXCLUDE — РАСШИРЕННАЯ", level=2)
doc.add_paragraph(
    "Базовая твоя EXCLUDE-строка + дополнения по шуму, который я увидел в топ-20 свежих:"
)
doc.add_paragraph(
    "Добавлено: изоляционн, пирометр, розетк, гигрометр, бытов++термометр, лабораторн++термометр, "
    "водонагреват, двер++металл, термошкаф, Тахион, интерфейс++защит, термографическ, "
    "термобель, трекинг, ботинк, КИПиА, сувенирн, термоформов, горелк++крематор, газоочистк, "
    "термическ++анализ, Netzsch, термоустойчив, термоэлектрич, термоплён/термоплен, "
    "бан++водян (бытовая баня), пиролитическ++ТМЦ, ОМ++МК, заглушк, термопар, термос++бытов."
).italic = True
codebox(doc, f"EXCLUDE — поле «По названию тендера». Длина: {len(EXC_01_NEW)} знаков.",
        EXC_01_NEW, "FCE4D6")

doc.add_heading("3.3. Прочие фильтры (UI)", level=2)
kv_table(doc, ["Фильтр", "Значение", "Комментарий"], [
    ["Поиск по полям", "ТОЛЬКО «Название тендера» (tender_name)",
     "Отключить «По документации» (tender_files)"],
    ["Дата публикации", "от = ВЧЕРА 00:00, до = пусто",
     "Раз в день обновлять «от» вручную (либо CLI каждый день создаёт временный автопоиск)"],
    ["НМЦК", "от = 4 000 000 ₽, до = пусто",
     "Приборы дешевле — не наш интерес"],
    ["Статус", "активные", "уже стоит"],
    ["Регион", "все", "не ограничиваем"],
    ["Тип закупки", "все (44-ФЗ, 223-ФЗ, ком., СНГ)", "не ограничиваем"],
    ["Сортировка", "tender_sysPublishDate.desc", "новые сверху"],
])

# 4. MDX_01_Sequencers
doc.add_heading("4. Автопоиск MDX_01_Sequencers (id 373417)", level=1)
doc.add_heading("4.1. INCLUDE (без изменений — твоя версия)", level=2)
codebox(doc, f"INCLUDE — поле «По названию тендера». Длина: {len(INC_MDX)} знаков.",
        INC_MDX, "E2EFDA")

doc.add_heading("4.2. EXCLUDE — РАСШИРЕННАЯ", level=2)
doc.add_paragraph(
    "Базовая твоя EXCLUDE (с ПЦР-блоком, single-cell и т.д.) + дополнения по шуму из топ-20:"
)
doc.add_paragraph(
    "Добавлено: дорожн++знак, ядерн, ДПО, повышен++квалификац, детск++площадк, ремонт++кровл, "
    "LECO++CS, сер++углерод++LECO, коленчат++вал, ремень++генератор, видеостен, экскурсионн, "
    "гирлянд, иллюминац, дизельн++топлив, электромонтажн, биохимическ++анализатор, "
    "автоматическ++биохим, IT++закупк, новогодн, праздничн++иллюминац, шкив++коленчат."
).italic = True
codebox(doc, f"EXCLUDE — поле «По названию тендера». Длина: {len(EXC_MDX_NEW)} знаков.",
        EXC_MDX_NEW, "FCE4D6")

doc.add_heading("4.3. Прочие фильтры (UI)", level=2)
kv_table(doc, ["Фильтр", "Значение", "Комментарий"], [
    ["Поиск по полям", "ТОЛЬКО «Название тендера» (tender_name)",
     "Отключить «По документации»"],
    ["Дата публикации", "от = ВЧЕРА 00:00, до = пусто", "обновлять «от» вручную ежедневно"],
    ["НМЦК", "от = 4 000 000 ₽, до = пусто",
     "Приборные секвенаторы дешевле 4 млн не бывают. (Для MDX_02 расходников — порог снимем.)"],
    ["Статус", "активные", "уже стоит"],
    ["Сортировка", "tender_sysPublishDate.desc", "новые сверху"],
])

# 5. Чеклист
doc.add_heading("5. Чек-лист пересоздания в UI", level=1)
steps = [
    "Открыть автопоиск 01_LC_LCMS_GPC_Prep (id 373343) → «Редактировать».",
    "В фильтре «Поиск по ключевым словам» — оставить «По названию тендера», снять галочку «По документации».",
    "В поле «Включать» — оставить как есть (INCLUDE не меняем).",
    "В поле «Исключать» — заменить на новую расширенную EXCLUDE-строку (раздел 3.2).",
    "Добавить фильтр «Дата публикации»: range «от = ВЧЕРА 00:00».",
    "Добавить фильтр «НМЦК»: range «от = 4 000 000 ₽».",
    "Сохранить.",
    "Открыть автопоиск MDX_01_Sequencers (id 373417) → повторить шаги 2-7 со строками из разделов 4.1–4.3.",
    "Сообщить мне — я через API замерю total_count и сделаю выборку топ-20 для оценки качества.",
]
for s in steps:
    doc.add_paragraph(s, style="List Number")

# 6. Tier-2 LLM
doc.add_heading("6. Tier-2 — LLM-проверка релевантности (проект)", level=1)
doc.add_paragraph(
    "Tier-1 (автопоиск в Tenderland) даёт грубое сито. После него — Tier-2 проверка каждого "
    "тендера через LLM, БЕЗ скачивания файлов. Дёшево и быстро."
)

doc.add_heading("6.1. На вход LLM подаётся", level=2)
for s in [
    "Имя автопоиска (например «01_LC_LCMS_GPC_Prep»).",
    "Краткое описание класса приборов («Жидкостная хроматография: ВЭЖХ/УВЭЖХ/ЖХ-МС/ГПХ/препаративная»).",
    "tender_regNumber, tender_name, tender_beginPrice, tender_lotCustomerShortName, "
    "tender_region, tender_typeName, tender_lotCategories, tender_endDate, tender_fileCount.",
]:
    doc.add_paragraph(s, style="List Bullet")

doc.add_heading("6.2. На выход LLM возвращает JSON", level=2)
codebox(doc, "Пример ответа LLM:",
        '{\n'
        '  "relevance": "pass" | "review" | "fail",\n'
        '  "confidence": 0.0..1.0,\n'
        '  "reasoning": "Краткое объяснение (1-2 предложения).",\n'
        '  "detected_class": "HPLC | LCMS-TQ | LCMS-QTOF | preparative | GPC | other",\n'
        '  "flags": ["maybe_consumables_only", "service_contract", "unclear_object"]\n'
        '}', "EDEDED")

doc.add_heading("6.3. Правила решения", level=2)
kv_table(doc, ["Решение", "Условие", "Что делаем"], [
    ["pass", "relevance=pass и confidence ≥ 0.7",
     "Скачиваем zip (Tier-3), отдаём анализатору"],
    ["review", "relevance=review ИЛИ pass с confidence 0.4-0.7",
     "Скачиваем zip, помечаем «требует ревью менеджера»"],
    ["fail", "relevance=fail",
     "Не скачиваем, в отчёт «отбраковано Tier-2», с причиной"],
])

doc.add_heading("6.4. Стоимость и модель", level=2)
doc.add_paragraph(
    "Модель — Claude Haiku (cheap) через LiteLLM, ~200 input + 80 output токенов на тендер. "
    "Стоимость ~$0.0001/тендер. На 50 тендеров в день по всем 21 автопоискам — ~$0.005/день, "
    "~$0.15/мес. Дешевле, чем кофе."
)
doc.add_paragraph(
    "Можно скормить пачкой по 10-20 тендеров в одном промпте для экономии (cache + batching) — "
    "тогда $0.05/мес."
)

doc.add_heading("6.5. Где живёт", level=2)
doc.add_paragraph(
    "Между Tier-1 (Searcher) и Tier-3 (file download). Реализуется в "
    "src/tenderland_bot/relevance/llm_filter.py. Тендеры с decision=fail НЕ тратят units на "
    "File/GetAll — это даёт самую большую экономию units (1 unit per file in archive)."
)

# 7. Дальнейшие шаги
doc.add_heading("7. Дальнейшие шаги", level=1)
doc.add_paragraph(
    "1) Ты пересоздаёшь 01_LC и MDX_01 по разделам 3–4 этого документа.\n"
    "2) Я через API замеряю новый total_count для обоих и беру топ-20 — оцениваем релевантность.\n"
    "3) Если шум упал до <20% (target для Tier-1) — раскатываю те же изменения на остальные 19 "
    "автопоисков (отдельный документ).\n"
    "4) Параллельно реализую Tier-2 LLM-фильтр (промпт, клиент LiteLLM, batch-режим) — на этих "
    "же двух пилотах прогоняем и проверяем качество.\n"
    "5) После Tier-1+Tier-2 на пилотах — переходим к Tier-3 (скачивание ТЗ) и Analyzer Module 1."
)

doc.add_paragraph(
    "После того как два пилота будут готовы — повторяю Word-документ для остальных 19 автопоисков "
    "с теми же расширениями EXCLUDE (соответствующим под каждую тему)."
)

doc.save(OUT)
print(f"Saved: {OUT}")
print(f"  EXCLUDE 01_LC:  {len(EXC_01)} -> {len(EXC_01_NEW)} (+{len(EXC_01_NEW) - len(EXC_01)})")
print(f"  EXCLUDE MDX_01: {len(EXC_MDX)} -> {len(EXC_MDX_NEW)} (+{len(EXC_MDX_NEW) - len(EXC_MDX)})")
